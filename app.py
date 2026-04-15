import streamlit as st
import os
import io
import re
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from crewai import Agent, Task, Crew, Process, LLM

# ===== 1. QUẢN LÝ TRẠNG THÁI =====
if 'final_srs' not in st.session_state: st.session_state.final_srs = ""
if 'final_testcases' not in st.session_state: st.session_state.final_testcases = ""

GROQ_MODELS = ["llama-3.3-70b-versatile", "llama-3.1-8b-instant", "openai/gpt-oss-120b", "openai/gpt-oss-20b"]

# KHUÔN MẪU BẢNG (Giữ nguyên độ gắt)
UC_TABLE_TEMPLATE = """
BẮT BUỘC TRÌNH BÀY ĐẶC TẢ DƯỚI DẠNG BẢNG MARKDOWN CÓ CẤU TRÚC SAU (Dùng thẻ <br> để xuống dòng):
| Thuộc tính | Chi tiết |
|---|---|
| **Mã UC / Tên** | UC-0X: [Tên Use Case] |
| **Tác nhân (Actor)** | [Ai thực hiện?] |
| **Tiền điều kiện** | [Điều kiện để bắt đầu] |
| **Luồng sự kiện chính (Main Flow)** | 1. ... <br> 2. ... <br> 3. ... |
| **Luồng rẽ nhánh / Lỗi (Alt Flow)** | - [Mã AF1]: ... <br> - [Mã AF2]: ... |
| **Hậu điều kiện** | [Trạng thái sau khi xong] |
| **Ràng buộc nghiệp vụ (Rules)** | [Quy tắc validate dữ liệu, định dạng...] |
"""

TESTCASE_TABLE_TEMPLATE = """
BẮT BUỘC TRÌNH BÀY DƯỚI DẠNG BẢNG MARKDOWN:
| TC-ID | Tên Test Case | Các bước thực hiện (Steps) | Dữ liệu đầu vào (Input) | Kết quả mong đợi (Expected) |
|---|---|---|---|---|
| TC-01 | ... | 1. ...<br>2. ... | ... | ... |
"""

# ===== 2. GIAO DIỆN CHÍNH =====
st.set_page_config(page_title="Ultimate BA/QA Architect", layout="wide")
st.title("🤖 ULTIMATE BA/QA: Hệ Thống Sinh Tài Liệu Tự Trị (Tự Động Mở Rộng)")

with st.sidebar:
    st.header("🔑 Cấu hình Hệ thống")
    st.session_state.api_key = st.text_input("Groq API Key", value=st.secrets.get("GROQ_API_KEY", ""), type="password")
    st.divider()
    m_strategist = st.selectbox("Product Strategist (Mở rộng ý tưởng)", GROQ_MODELS, index=0)
    m_ba = st.selectbox("Lead BA (Viết SRS)", GROQ_MODELS, index=0)
    m_qa = st.selectbox("Principal QA (Bắt lỗi)", GROQ_MODELS, index=2)
    m_tpm = st.selectbox("Tech Lead (Chốt sổ)", GROQ_MODELS, index=0)

user_idea = st.text_area("🚀 Nhập ý tưởng sơ khai (Hệ thống sẽ tự động Brainstorm & Mở rộng):", height=120)

# ===== 3. CORE LOGIC 1: AUTONOMOUS SRS (CÓ BƯỚC MỞ RỘNG) =====
if st.button("🔥 BƯỚC 1: TỰ ĐỘNG PHÂN TÍCH, MỞ RỘNG & TẠO SRS", type="primary"):
    if not st.session_state.api_key or not user_idea: 
        st.error("Bác nhập thiếu Key hoặc Ý tưởng rồi!"); st.stop()
    
    os.environ["GROQ_API_KEY"] = st.session_state.api_key
    st.session_state.final_srs = ""
    st.session_state.final_testcases = ""

    # Khởi tạo 4 bộ não
    llm_strategist = LLM(model=f"groq/{m_strategist}", temperature=0.7) # Nhiệt độ cao để sáng tạo
    llm_ba = LLM(model=f"groq/{m_ba}", temperature=0.3)
    llm_qa = LLM(model=f"groq/{m_qa}", temperature=0.1)
    llm_tpm = LLM(model=f"groq/{m_tpm}", temperature=0.1)

    agent_strategist = Agent(
        role="Product Strategist & Consultant", 
        goal="Đọc ý tưởng sơ khai, bóc tách và SÁNG TẠO MỞ RỘNG thành một danh sách Use Case đồ sộ, bao phủ mọi ngóc ngách của hệ thống.", 
        backstory="Chuyên gia tư vấn chiến lược sản phẩm (CPO). Bạn luôn nhìn thấy những tính năng ẩn mà khách hàng chưa nghĩ tới (như quản lý thanh toán, admin, bảo mật, gửi mail...).", 
        llm=llm_strategist
    )
    agent_ba = Agent(role="Lead BA", goal="Viết đặc tả sâu sắc, KẺ BẢNG USE CASE.", backstory="BA 15 năm kinh nghiệm.", llm=llm_ba)
    agent_qa = Agent(role="Principal QA", goal="Tìm lỗ hổng logic.", backstory="Chuyên gia bắt lỗi.", llm=llm_qa)
    agent_tpm = Agent(role="Tech Lead", goal="Chốt hạ tài liệu.", backstory="Người ra quyết định.", llm=llm_tpm)

    with st.status("🚀 Hệ thống đang tư duy chiến lược...", expanded=True) as status:
        full_document = ""
        
        # --- BƯỚC 1: TỔNG QUAN ---
        st.write("⚙️ Đang lập Outline & Tổng quan...")
        t_ch1_draft = Task(description=f"Dựa vào ý tưởng: '{user_idea}'. Viết CHƯƠNG 1: TỔNG QUAN. Gồm: Mục tiêu, Phạm vi, Đối tượng.", expected_output="Bản nháp", agent=agent_ba)
        t_ch1_final = Task(description="Chuẩn hóa bản nháp Chương 1.", expected_output="Bản chốt", agent=agent_tpm)
        crew_ch1 = Crew(agents=[agent_ba, agent_tpm], tasks=[t_ch1_draft, t_ch1_final], verbose=False)
        full_document += "CHƯƠNG 1: TỔNG QUAN DỰ ÁN\n" + getattr(crew_ch1.kickoff(), 'raw', "") + "\n\n"

        # --- BƯỚC 2: BRAINSTORM & MỞ RỘNG USE CASE (BẢN FIX LỖI PARSE) ---
        time.sleep(2)
        st.write("💡 AI Strategist đang mở rộng ý tưởng và vạch ra toàn bộ Use Case...")
        t_expand_uc = Task(
            description=f"""Ý tưởng gốc của khách hàng: '{user_idea}'.
            Nhiệm vụ: Hãy đóng vai CPO, vạch ra DANH SÁCH TOÀN BỘ Use Case cần thiết. Chia thành các phân hệ rõ ràng (Ví dụ: Phân hệ User, Admin...).
            
            ⚠️ LỆNH BẮT BUỘC ⚠️: 
            Mỗi Use Case bạn nghĩ ra PHẢI nằm trên 1 dòng riêng biệt và BẮT BUỘC bắt đầu bằng tiền tố "[UC]: ".
            Ví dụ:
            Phân hệ Khách hàng:
            [UC]: Đăng nhập hệ thống qua Gmail
            [UC]: Quét mã QR truy cập bia mộ
            
            Tuyệt đối không được quên tiền tố [UC]: trước mỗi chức năng.""",
            expected_output="Danh sách Use Case có tiền tố [UC]:", 
            agent=agent_strategist
        )
        crew_expand = Crew(agents=[agent_strategist], tasks=[t_expand_uc], verbose=False)
        
        # Bắt lỗi an toàn khi lấy dữ liệu từ CrewAI
        try:
            res_expand = crew_expand.kickoff()
            uc_raw_list = getattr(res_expand, 'raw', str(res_expand))
        except Exception as e:
            uc_raw_list = str(e)
            st.warning("Có chút lỗi khi nhận dữ liệu từ AI, hệ thống đang tự động khôi phục...")
        
        # Parse danh sách chuẩn xác 100% dựa vào Key "[UC]:"
        all_use_cases = []
        for line in uc_raw_list.split('\n'):
            if '[UC]:' in line:
                # Cắt lấy phần tên đằng sau chữ [UC]: và dọn dẹp ký tự thừa
                uc_name = line.split('[UC]:')[-1].replace('**', '').replace('*', '').strip()
                if len(uc_name) > 3:
                    all_use_cases.append(uc_name)
        
        # Fallback (Phương án dự phòng): Nếu AI cãi lệnh không sinh chữ [UC]:
        if not all_use_cases:
            import re
            # Vét tất cả các dòng có đánh số (1., 2.) hoặc gạch đầu dòng (-, *)
            all_use_cases = [re.sub(r'^[\d\.\-\*\s]+', '', line).strip() for line in uc_raw_list.split('\n') if re.match(r'^[\s]*[\-\*\d]', line) and len(line) > 5]
            
        # Fallback cuối cùng nếu có lỗi mạng API (chống crash app)
        if not all_use_cases: 
            all_use_cases = ["Quản lý hồ sơ người đã khuất", "Tương tác dâng hương ảo", "Thanh toán gói cước 6 tháng/1 năm", "Tích hợp và đăng nhập Gmail"]
        
        # Lọc ra 8 Use Case quan trọng nhất để chạy chi tiết (Tránh Rate Limit)
        use_cases_to_run = all_use_cases[:8] if len(all_use_cases) > 8 else all_use_cases
        
        st.success(f"🧠 Strategist đã nghĩ ra {len(all_use_cases)} Use Case! Đang tiến hành phân tích sâu {len(use_cases_to_run)} Use Case cốt lõi nhất:")
        for uc in use_cases_to_run: 
            st.markdown(f"- {uc}")
        
        full_document += "CHƯƠNG 2: ĐẶC TẢ USE CASE CHI TIẾT\n"

        # --- BƯỚC 3: VÒNG LẶP KẺ BẢNG TỪNG USE CASE ---
        for idx, uc_name in enumerate(use_cases_to_run):
            st.write(f"🔄 Đang mổ xẻ UC {idx+1}/{len(use_cases_to_run)}: **{uc_name}**")
            t_uc_draft = Task(description=f"Đặc tả UC: '{uc_name}'.\n\n{UC_TABLE_TEMPLATE}", expected_output="Bảng nháp UC", agent=agent_ba)
            t_uc_critic = Task(description=f"Soi bảng nháp UC '{uc_name}'. Tìm lỗ hổng: API sập? Spam click? Data rỗng?", expected_output="Lỗi UC", agent=agent_qa)
            t_uc_final = Task(description=f"Cập nhật lỗi vào bảng đặc tả UC '{uc_name}'. BẮT BUỘC GIỮ NGUYÊN FORMAT BẢNG MARKDOWN.", expected_output="Bảng chốt UC", agent=agent_tpm)
            
            crew_uc = Crew(agents=[agent_ba, agent_qa, agent_tpm], tasks=[t_uc_draft, t_uc_critic, t_uc_final], verbose=False)
            try:
                time.sleep(3)
                uc_result = getattr(crew_uc.kickoff(), 'raw', "")
            except Exception:
                st.warning("⏳ Rate Limit! Tạm nghỉ 12s...")
                time.sleep(12)
                uc_result = getattr(crew_uc.kickoff(), 'raw', "")
            
            full_document += f"\n### 2.{idx+1}. Use Case: {uc_name}\n" + uc_result + "\n"

        # --- BƯỚC 4: YÊU CẦU PHI CHỨC NĂNG ---
        time.sleep(3)
        st.write("⚙️ Viết yêu cầu phi chức năng...")
        t_ch3_draft = Task(description=f"Viết CHƯƠNG 3: YÊU CẦU PHI CHỨC NĂNG cho: {user_idea}.", expected_output="Nháp", agent=agent_ba)
        t_ch3_final = Task(description="Chuẩn hóa bản nháp Chương 3", expected_output="Chốt", agent=agent_tpm)
        crew_ch3 = Crew(agents=[agent_ba, agent_tpm], tasks=[t_ch3_draft, t_ch3_final], verbose=False)
        full_document += "\nCHƯƠNG 3: YÊU CẦU PHI CHỨC NĂNG & DỮ LIỆU\n" + getattr(crew_ch3.kickoff(), 'raw', "")
        
        st.session_state.final_srs = full_document
        status.update(label="✅ ĐÃ HOÀN THÀNH SRS DẠNG BẢNG (CÓ MỞ RỘNG)!", state="complete")

# ===== 4. CORE LOGIC 2: KHU VỰC QA (TẠO TEST CASE) =====
if st.session_state.final_srs:
    st.divider()
    st.markdown("### 🧪 KHU VỰC KIỂM THỬ (QA AREA)")
    
    if st.button("🚀 BƯỚC 2: TỰ ĐỘNG SINH TEST CASE DẠNG BẢNG", type="primary"):
        os.environ["GROQ_API_KEY"] = st.session_state.api_key
        # Khởi tạo lại não QA để tách biệt luồng
        llm_qa = LLM(model=f"groq/{m_qa}", temperature=0.1)
        agent_qa_test = Agent(
            role="Senior QA Automation", 
            goal="Kẻ bảng Test Case chuẩn ISTQB.", 
            backstory="Trùm kiểm thử, không để lọt bất kỳ case ngoại lệ nào.", 
            llm=llm_qa
        )

        t_testcase = Task(
            description=f"Đọc tài liệu SRS sau:\n{st.session_state.final_srs[:4000]}\n\n{TESTCASE_TABLE_TEMPLATE}",
            expected_output="Bảng Test Case", 
            agent=agent_qa_test
        )
        crew_tc = Crew(agents=[agent_qa_test], tasks=[t_testcase], verbose=False)

        with st.status("🔬 QA đang phân tích SRS và kẻ bảng Test Case...", expanded=True) as status:
            try:
                st.session_state.final_testcases = getattr(crew_tc.kickoff(), 'raw', "")
                status.update(label="✅ ĐÃ HOÀN TẤT SINH BỘ TEST CASE DẠNG BẢNG!", state="complete")
            except Exception as e:
                st.error(f"Lỗi hệ thống: {e}")
                status.update(label="❌ Có lỗi xảy ra!", state="error")

# ===== 5. BỘ MÁY PARSE MARKDOWN SANG WORD TABLE & XUẤT FILE =====
if st.session_state.final_srs:
    st.divider()
    tab1, tab2 = st.tabs(["📄 TÀI LIỆU SRS (MASTER)", "🧪 BỘ TEST CASE (QA)"])
    
    with tab1: 
        st.text_area("Nội dung SRS (Markdown)", st.session_state.final_srs, height=500)
    with tab2: 
        if st.session_state.final_testcases: 
            st.text_area("Nội dung Test Case (Markdown)", st.session_state.final_testcases, height=500)
        else:
            st.warning("⚠️ Bác chưa bấm nút 'Tạo Test Case' ở trên.")

    def create_export_docx():
        doc = Document()
        doc.sections[0].left_margin = Cm(2.5)
        doc.sections[0].right_margin = Cm(2.5)
        
        # Trang bìa
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n\n\nENTERPRISE SRS & TEST PLAN\n\n")
        run.bold = True
        run.font.size = Pt(20)
        
        # Gộp toàn bộ văn bản
        final_text = st.session_state.final_srs
        if st.session_state.final_testcases:
            final_text += "\n\nCHƯƠNG 4: KỊCH BẢN KIỂM THỬ (TEST SUITE)\n" + st.session_state.final_testcases

        lines = final_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # HỆ THỐNG DỊCH BẢNG MARKDOWN THÀNH BẢNG WORD
            if line.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                if len(table_lines) > 2: # Chắc chắn có header, dòng kẻ, và data
                    # Lấy Header
                    headers = [col.strip() for col in table_lines[0].split('|')[1:-1]]
                    if headers:
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid' # Đóng khung viền đen chuẩn Word
                        
                        # Điền Header
                        for j, h in enumerate(headers):
                            table.rows[0].cells[j].text = h.replace('**', '')
                            table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
                        
                        # Điền Data (Bỏ qua dòng thứ 2 là dòng chứa ---|---)
                        for t_line in table_lines[2:]:
                            cols = [col.strip() for col in t_line.split('|')[1:-1]]
                            row_cells = table.add_row().cells
                            for j, c in enumerate(cols):
                                if j < len(row_cells):
                                    # Chuyển đổi <br> thành xuống dòng thực sự trong Word
                                    clean_text = c.replace('<br>', '\n').replace('**', '')
                                    row_cells[j].text = clean_text
                continue # Xong khối bảng thì đi tiếp, không in text thường nữa

            # Xử lý Heading và Text bình thường
            if re.match(r'^#+', line):
                level = line.count('#')
                heading_text = line.replace('#', '').strip()
                doc.add_heading(heading_text, level=level if level <= 9 else 9)
            elif re.match(r'^CHƯƠNG \d+', line) or re.match(r'^2\.\d+\. Use Case', line):
                h = doc.add_heading(line, level=1 if 'CHƯƠNG' in line else 2)
                for r in h.runs: 
                    r.font.color.rgb = RGBColor(0,0,0)
            else:
                p = doc.add_paragraph(line)
                if p.runs: 
                    p.runs[0].font.name = "Times New Roman"
            i += 1
        
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    st.download_button(
        "📥 TẢI XUỐNG FILE WORD TỔNG HỢP (CÓ KẺ BẢNG CHUẨN)", 
        create_export_docx(), 
        "Enterprise_SRS_Master.docx", 
        type="primary", 
        use_container_width=True
    )